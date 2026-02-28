import streamlit as st
import pandas as pd
import sqlite3
import os
import hashlib
import time
from supabase import create_client, Client
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
from PIL import Image

# --- 1. CONFIGURAÇÃO COM ÍCONE DE INSTALAÇÃO ---
st.set_page_config(page_title="AEE Conecta", layout="centered", page_icon="logo.png")

# Esse bloco reforça para o Windows/Android/iPhone qual imagem usar no ícone de atalho
st.markdown(
    """
    <head>
        <meta name="google" content="notranslate">
        <link rel="icon" href="https://raw.githubusercontent.com/marcoshgomes/aee-conecta/main/logo.png">
        <link rel="apple-touch-icon" href="https://raw.githubusercontent.com/marcoshgomes/aee-conecta/main/logo.png">
        <link rel="shortcut icon" href="https://raw.githubusercontent.com/marcoshgomes/aee-conecta/main/logo.png">
    </head>
    <script>
        document.documentElement.lang = 'pt-br';
        document.documentElement.classList.add('notranslate');
    </script>
    """,
    unsafe_allow_html=True
)

# --- 2. CONEXÃO SUPABASE ---
try:
    URL = st.secrets["supabase"]["url"]
    KEY = st.secrets["supabase"]["key"]
    supabase: Client = create_client(URL, KEY)
except:
    st.error("Erro nas credenciais do Supabase no arquivo secrets.toml.")
    st.stop()

# Garantir pastas locais
if not os.path.exists("fotos_alunos"): os.makedirs("fotos_alunos")

# --- 3. FUNÇÕES DE APOIO ---
def hash_pw(senha):
    return hashlib.sha256(str(senha).encode()).hexdigest()

def registrar_log(rf):
    try:
        supabase.table("logs").insert({"rf": rf, "data_hora": datetime.now().strftime('%d/%m/%Y %H:%M:%S')}).execute()
    except:
        pass

def load_professores():
    try:
        res = supabase.table("professores").select("*").execute()
        return pd.DataFrame(res.data)
    except:
        return pd.DataFrame()

def load_estudantes():
    try:
        res = supabase.table("estudantes").select("*").execute()
        return pd.DataFrame(res.data)
    except:
        return pd.DataFrame()

# --- 4. FUNÇÕES DE GERAÇÃO DE WORD ---
def gerar_folha_rosto(dados):
    doc = Document()
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("CEU EMEF Prof.ª MARA CRISTINA TARTAGLIA SENA\nAEE - ATENDIMENTO EDUCACIONAL ESPECIALIZADO\nREGISTRO - ATIVIDADE FLEXIBILIZADA").bold = True
    doc.add_paragraph(f"ANO LETIVO {datetime.now().year}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if dados.get('foto_path'):
        try:
            res_foto = supabase.storage.from_("fotos_perfil").download(dados['foto_path'])
            doc.add_picture(BytesIO(res_foto), width=Inches(2.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    t = doc.add_table(rows=4, cols=1); t.style = 'Table Grid'
    t.rows[0].cells[0].text = f"ESTUDANTE: {dados.get('aluno', 'N/A')}"
    t.rows[1].cells[0].text = f"TURMA: {dados.get('turma', 'N/A')}"
    col_nec = next((x for x in list(dados.index) if "nec" in x), "necessidades")
    t.rows[2].cells[0].text = f"DEFICIÊNCIA/CONDIÇÃO: {dados.get(col_nec, 'N/A')}"
    t.rows[3].cells[0].text = f"DATA DE NASCIMENTO: {dados.get('data_nascimento', 'N/A')}"
    
    doc.add_heading('PERFIL E OBSERVAÇÕES DO PROFESSOR PAEE:', level=3)
    doc.add_paragraph(str(dados.get('observacoes_gerais', '')))
    
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

def gerar_relatorio_aula(df_rels, nome_aluno, turma_aluno, df_professores):
    doc = Document()
    for i, row in df_rels.iterrows():
        rf_aula = row['rf_professor']
        filtro_p = df_professores[df_professores['rf'] == rf_aula]
        nome_p = filtro_p.iloc[0]['nome'] if not filtro_p.empty else "Professor não identificado"

        h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.add_run("CEU EMEF Prof.ª MARA CRISTINA TARTAGLIA SENA\nAEE - ATENDIMENTO EDUCACIONAL ESPECIALIZADO\nREGISTRO - ATIVIDADE FLEXIBILIZADA").bold = True
        p_title = doc.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.add_run(f"{row['bimestre']} – ANO LETIVO {datetime.now().year}").bold = True
        
        doc.add_paragraph(f"ESTUDANTE: {nome_aluno}\nTURMA: {turma_aluno}")
        doc.add_paragraph(f"PROFESSOR: {nome_p} | DISCIPLINA/TEMA: {row.get('disciplina_tema', 'N/A')}")
        
        p_sim = "x" if row.get('participou_aula') == "Sim" else " "
        p_nao = "x" if row.get('participou_aula') == "Não" else " "
        txt_part = doc.add_paragraph(f"O ESTUDANTE PARTICIPOU DA SUA AULA? ( {p_sim} ) SIM ( {p_nao} ) NÃO.")
        if row.get('participou_aula') == "Não":
            txt_part.add_run(f" RELATE O MOTIVO: {row.get('motivo_nao_participou', '')}")
            
        doc.add_heading('ATIVIDADES PLANEJADAS:', level=3); doc.add_paragraph(str(row['planejado']))
        doc.add_heading('ATIVIDADE REALIZADA COM O ESTUDANTE:', level=3); doc.add_paragraph(str(row['realizado']))
        
        if row['foto_path']:
            try:
                res_foto = supabase.storage.from_("fotos_aee").download(row['foto_path'])
                doc.add_picture(BytesIO(res_foto), width=Inches(3.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: pass
        
        doc.add_heading('COMO FOI A PARTICIPAÇÃO DO ESTUDANTE?', level=3)
        parts = str(row['participacao']).split(", ")
        opcoes = ["REALIZOU COM AUTONOMIA", "REALIZOU COM APOIO E INTERVENÇÃO DE UM ADULTO", "REALIZOU WITH APOIO DE UM COLEGA", "NÃO REALIZOU"]
        for op in opcoes:
            check = "x" if op in parts else " "
            doc.add_paragraph(f"( {check} ) {op}")
        
        doc.add_paragraph(f"\nDATA DE REALIZAÇÃO DA ATIVIDADE: {row['data']}")
        if i < len(df_rels) - 1: doc.add_page_break()
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

# --- 5. LÓGICA DE LOGIN ---
df_prof = load_professores()

if "logged_in" not in st.session_state: st.session_state.logged_in = False
if "change_pw" not in st.session_state: st.session_state.change_pw = False

if not st.session_state.logged_in:
    st.title("💠 AEE Conecta - Login")
    rf_in = st.text_input("RF").strip()
    pw_in = st.text_input("Senha", type="password").strip()

    if st.button("Entrar"):
        if df_prof.empty:
            supabase.table("professores").insert({"rf": rf_in, "nome": "Gestor Mestre", "perfil": "gestao"}).execute()
            st.rerun()
        
        user_db = df_prof[df_prof['rf'] == rf_in]
        if not user_db.empty:
            res_creds = supabase.table("credenciais").select("senha_hash").eq("rf", rf_in).execute()
            
            if not res_creds.data:
                if pw_in == rf_in:
                    st.session_state.change_pw, st.session_state.temp_rf = True, rf_in
                    st.rerun()
                else:
                    st.warning("Primeiro acesso? Use seu RF como senha.")
            else:
                if hash_pw(pw_in) == res_creds.data[0]['senha_hash']:
                    st.session_state.logged_in, st.session_state.u_rf = True, rf_in
                    st.session_state.u_nome = user_db.iloc[0]['nome']
                    st.session_state.u_perfil = str(user_db.iloc[0]['perfil']).lower().strip().replace('çã', 'ca')
                    registrar_log(rf_in); st.rerun()
                else:
                    st.error("Senha incorreta.")
        else:
            st.error("Usuário não cadastrado.")

    if st.session_state.change_pw:
        st.divider()
        st.warning("⚠️ CADASTRE UMA SENHA PESSOAL")
        n_pw = st.text_input("Nova Senha (min. 6 carac.)", type="password")
        c_pw = st.text_input("Confirme a Senha", type="password")
        if st.button("Salvar Nova Senha"):
            if len(n_pw) >= 6 and n_pw == c_pw:
                supabase.table("credenciais").insert({"rf": st.session_state.temp_rf, "senha_hash": hash_pw(n_pw)}).execute()
                st.session_state.change_pw = False
                st.toast("✅ Senha cadastrada!")
                time.sleep(1.5); st.rerun()
            else:
                st.error("Verifique os dados.")

else:
    # --- 6. INTERFACE LOGADA ---
    df_alunos = load_estudantes()
    st.sidebar.title(f"Olá, {st.session_state.u_nome}")
    menu = st.sidebar.radio("Navegação", ["Início", "Lançar Relatório", "Painel de Documentos", "Sair"])
    if menu == "Sair": st.session_state.logged_in = False; st.rerun()
    super_perfis = ["gestao", "gestor", "paee", "direcao", "coordenador"]

    # --- INÍCIO ---
    if menu == "Início":
        st.title("🏠 Bem-vindo ao AEE Conecta")
        st.markdown(f"### Olá, {st.session_state.u_nome}!")
        st.divider()
        st.info("""
        Este sistema foi desenvolvido para simplificar o registro e a gestão dos atendimentos do AEE.
        
        **Como utilizar:**
        *   **Professores:** Utilize o menu ao lado e clique em **Lançar Relatório** para registrar suas aulas.
        *   **Gestores/PAEE:** Acesse o **Painel de Documentos** para baixar Folhas de Rosto e gerenciar perfis.
        
        *Dica: No celular, você pode usar o microfone do teclado para ditar os textos das atividades!*
        """)
        st.success("Selecione uma opção no menu lateral para começar.")
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        st.markdown(
            """
            <div style='text-align: center; color: #888888; font-size: 0.9em; border-top: 1px solid #eeeeee; padding-top: 20px;'>
                <b>By Prof. Marcão</b><br>
                Software de apoio pedagógico - Freeware
            </div>
            """,
            unsafe_allow_html=True
        )

    # --- LANÇAR RELATÓRIO ---
    elif menu == "Lançar Relatório":
        st.header("📝 Lançar Relatório de Aula")
        
        # Inicialização da memória de replicação
        if "tema_val" not in st.session_state: st.session_state.tema_val = ""
        if "plan_val" not in st.session_state: st.session_state.plan_val = ""
        if "form_reset_key" not in st.session_state: st.session_state.form_reset_key = 0

        if df_alunos.empty:
            st.warning("Aguardando cadastro de alunos pela Gestão/PAEE.")
        else:
            # 1. FILTRO POR TURMA
            lista_turmas = ["Todas"] + sorted(df_alunos['turma'].unique().tolist())
            turma_sel = st.selectbox("1. Filtrar por Turma:", lista_turmas)
            
            df_f = df_alunos[df_alunos['turma'] == turma_sel].copy() if turma_sel != "Todas" else df_alunos.copy()
            df_f['exibicao'] = df_f['aluno'] + " - " + df_f['turma']
            
            # 2. SELEÇÃO DO ESTUDANTE
            lista_est = ["Selecione o Estudante..."] + sorted(df_f['exibicao'].tolist())
            al_sel_visual = st.selectbox("2. Escolha o aluno:", lista_est, key=f"al_sel_{st.session_state.form_reset_key}")

            if al_sel_visual != "Selecione o Estudante...":
                nome_puro = al_sel_visual.split(" - ")[0]
                al_inf = df_alunos[df_alunos['aluno'] == nome_puro].iloc[0]
                
                with st.container(border=True):
                    # --- NOVO: IDENTIFICAÇÃO VISUAL DO ALUNO ---
                    col_foto, col_info = st.columns([1, 4])
                    
                    with col_foto:
                        if al_inf.get('foto_path'):
                            try:
                                # Busca a foto oficial do perfil (fotos_perfil)
                                res_foto = supabase.storage.from_("fotos_perfil").download(al_inf['foto_path'])
                                st.image(BytesIO(res_foto), width=100) # Tamanho pequeno mas visível
                            except:
                                st.write("🖼️ (Erro na foto)")
                        else:
                            st.write("🖼️ (Sem foto)")
                    
                    with col_info:
                        st.subheader(nome_puro)
                        st.write(f"**Turma:** {al_inf['turma']} | **Registro:** {al_inf['registro']}")
                        st.caption(f"**Condição:** {al_inf['necessidades']}")
                    
                    st.divider()
                    
                    # --- DADOS DA AULA ---
                    col1, col2 = st.columns(2)
                    with col1:
                        dt = st.date_input("Data da Atividade", datetime.now())
                    with col2:
                        bm = st.selectbox("Bimestre", ["1º Bimestre", "2º Bimestre", "3º Bimestre", "4º Bimestre"])
                    
                    tm = st.text_input("Disciplina ou Tema da Aula", value=st.session_state.tema_val)
                    pl = st.text_area("Atividades Planejadas", value=st.session_state.plan_val)
                    
                    st.divider()
                    st.subheader("📝 Parecer Individual")
                    
                    p_a = st.radio("O estudante participou?", ["Sim", "Não"], horizontal=True, key=f"pa_{st.session_state.form_reset_key}")
                    
                    if p_a == "Não":
                        mot = st.text_area("Relate o motivo da não participação:", key=f"mot_{st.session_state.form_reset_key}")
                        re = ""
                        pn = []
                    else:
                        mot = ""
                        re = st.text_area("Atividades Realizadas (Desempenho individual)", key=f"re_{st.session_state.form_reset_key}")
                        pn = st.multiselect("Nível de Participação:", 
                                           ["REALIZOU COM AUTONOMIA", "APOIO ADULTO", "APOIO COLEGA", "NÃO REALIZOU"],
                                           key=f"pn_{st.session_state.form_reset_key}")
                    
                    ft = st.file_uploader("Anexar foto do registro diário", type=['png', 'jpg', 'jpeg'], key=f"ft_{st.session_state.form_reset_key}")
                    
                    st.divider()
                    replicar = st.checkbox("Manter 'Tema' e 'Planejado' para o próximo registro?", value=True)

                    if st.button("💾 Salvar Relatório Individual"):
                        if not tm or not pl:
                            st.error("Preencha o Tema e o Planejamento.")
                        else:
                            p_f = ""
                            if ft:
                                p_f = f"aula_{datetime.now().strftime('%Y%m%d%H%M%S')}_{al_inf['registro']}.png"
                                supabase.storage.from_("fotos_aee").upload(p_f, ft.getvalue())
                            
                            supabase.table("relatorios").insert({
                                "data": dt.strftime('%d/%m/%Y'), "rf_professor": st.session_state.u_rf, 
                                "registro_aluno": str(al_inf['registro']), "bimestre": bm, "participou_aula": p_a, 
                                "motivo_nao_participou": mot, "disciplina_tema": tm, "planejado": pl, 
                                "realizado": re, "participacao": ", ".join(pn), "foto_path": p_f
                            }).execute()
                            
                            if replicar:
                                st.session_state.tema_val, st.session_state.plan_val = tm, pl
                            else:
                                st.session_state.tema_val, st.session_state.plan_val = "", ""
                            
                            st.toast(f"✅ Registro de {nome_puro} salvo!")
                            time.sleep(1)
                            st.session_state.form_reset_key += 1
                            st.rerun()

    # --- PAINEL DE DOCUMENTOS ---
    elif menu == "Painel de Documentos":
        st.title("📂 Painel de Documentos")
        super_perfis = ["gestao", "gestor", "paee", "direcao", "coordenador"]
        
        # 1. Definição Dinâmica das Abas (Professor agora vê 'Meus Registros')
        list_tabs = ["📄 Documentos", "✏️ Alterar ou Excluir"]
        if st.session_state.u_perfil in super_perfis:
            list_tabs += ["👤 Gestão de Alunos", "👥 Gestão de Professores", "🔒 Segurança e Reset"]
        
        abas = st.tabs(list_tabs)

        # --- ABA 0: DOWNLOAD DE DOCUMENTOS (Impressão) ---
        with abas[0]:
            if df_alunos.empty:
                st.info("Nenhum aluno cadastrado no sistema.")
            else:
                df_alunos['exibicao_imp'] = df_alunos['aluno'] + " - " + df_alunos['turma']
                al_f_v = st.selectbox("Selecione o Aluno para Documentos:", sorted(df_alunos['exibicao_imp'].tolist()), key="sel_doc_imp")
                al_f_nome = al_f_v.split(" - ")[0]
                d_f = df_alunos[df_alunos['aluno'] == al_f_nome].iloc[0]
                
                c1, c2 = st.columns(2)
                c1.download_button("📥 Baixar Folha de Rosto", gerar_folha_rosto(d_f), f"Rosto_{al_f_nome}.docx")
                
                # Regra de Visualização: Gestão vê tudo, Professor só vê o dele
                query = supabase.table("relatorios").select("*").eq("registro_aluno", str(d_f['registro']))
                if st.session_state.u_perfil not in super_perfis:
                    query = query.eq("rf_professor", st.session_state.u_rf)
                
                res = query.execute()
                df_res = pd.DataFrame(res.data)
                
                if not df_res.empty:
                    bim_f = st.selectbox("Filtrar Bimestre para Impressão:", ["Todos", "1º Bimestre", "2º Bimestre", "3º Bimestre", "4º Bimestre"])
                    if bim_f != "Todos": df_res = df_res[df_res['bimestre'] == bim_f]
                    if not df_res.empty:
                        c2.download_button(f"📥 Baixar Relatórios ({len(df_res)})", gerar_relatorio_aula(df_res, al_f_nome, d_f['turma'], df_prof), f"Relatos_{al_f_nome}.docx")
                else:
                    st.warning("Sem relatórios disponíveis para impressão conforme seu perfil.")

        # --- ABA 1: ALTERAR OU EXCLUIR (NOVA FUNCIONALIDADE) ---
        with abas[1]:
            st.subheader("Gerenciar Meus Registros")
            st.write("Aqui você pode corrigir erros ou excluir aulas lançadas por você.")
            
            al_ed_v = st.selectbox("Selecione o Aluno para ver seus relatórios:", ["Selecione..."] + sorted(df_alunos['exibicao_imp'].tolist()), key="sel_edit_rel")
            
            if al_ed_v != "Selecione...":
                nome_p_ed = al_ed_v.split(" - ")[0]
                al_inf_ed = df_alunos[df_alunos['aluno'] == nome_p_ed].iloc[0]
                
                # Busca relatórios que o usuário logado PODE editar
                q_ed = supabase.table("relatorios").select("*").eq("registro_aluno", str(al_inf_ed['registro']))
                if st.session_state.u_perfil not in super_perfis:
                    q_ed = q_ed.eq("rf_professor", st.session_state.u_rf)
                
                df_ed = pd.DataFrame(q_ed.execute().data)
                
                if df_ed.empty:
                    st.info("Nenhum registro seu encontrado para este aluno.")
                else:
                    df_ed['label'] = df_ed['data'] + " - " + df_ed['disciplina_tema']
                    rel_sel_label = st.selectbox("Escolha o registro para alterar:", df_ed['label'].tolist())
                    rel_data = df_ed[df_ed['label'] == rel_sel_label].iloc[0]
                    
                    with st.form("form_correcao_aula"):
                        st.warning(f"Modo Edição: Aula de {rel_data['data']}")
                        new_tm = st.text_input("Tema/Disciplina", value=rel_data['disciplina_tema'])
                        new_pl = st.text_area("Atividades Planejadas", value=rel_data['planejado'])
                        new_re = st.text_area("Atividades Realizadas", value=rel_data['realizado'])
                        new_pa = st.radio("Participou?", ["Sim", "Não"], index=0 if rel_data['participou_aula'] == "Sim" else 1)
                        new_pn = st.multiselect("Nível:", ["REALIZOU COM AUTONOMIA", "APOIO ADULTO", "APOIO COLEGA", "NÃO REALIZOU"], default=str(rel_data['participacao']).split(", "))
                        new_ft = st.file_uploader("Substituir foto (opcional)")
                        
                        b1, b2 = st.columns(2)
                        if b1.form_submit_button("💾 Salvar Alterações"):
                            p_f_update = rel_data['foto_path']
                            if new_ft:
                                p_f_update = f"aula_{datetime.now().strftime('%Y%m%d%H%M%S')}.png"
                                supabase.storage.from_("fotos_aee").upload(p_f_update, new_ft.getvalue())
                            
                            supabase.table("relatorios").update({
                                "disciplina_tema": new_tm, "planejado": new_pl, "realizado": new_re,
                                "participou_aula": new_pa, "participacao": ", ".join(new_pn), "foto_path": p_f_update
                            }).eq("id", rel_data['id']).execute()
                            st.toast("✅ Registro atualizado!"); time.sleep(1); st.rerun()
                            
                        if b2.form_submit_button("❌ EXCLUIR DEFINITIVAMENTE"):
                            if rel_data['foto_path']:
                                try: supabase.storage.from_("fotos_aee").remove([rel_data['foto_path']])
                                except: pass
                            supabase.table("relatorios").delete().eq("id", rel_data['id']).execute()
                            st.toast("⚠️ Registro removido!"); time.sleep(1); st.rerun()

        # --- ABAS DE GESTÃO (SÓ APARECEM PARA SUPER_PERFIS) ---
        if st.session_state.u_perfil in super_perfis:
            with abas[2]: # GESTÃO DE ALUNOS
                st.subheader("Gestão de Alunos")
                if "al_form_id" not in st.session_state: st.session_state.al_form_id = 0
                modo_a = st.radio("Ação Estudante:", ["Novo", "Editar/Excluir"], horizontal=True, key=f"ma_{st.session_state.al_form_id}")
                al_edit = None
                if modo_a == "Editar/Excluir" and not df_alunos.empty:
                    df_alunos['ex_g'] = df_alunos['aluno'] + " - " + df_alunos['turma']
                    sel_a_g = st.selectbox("Escolha:", sorted(df_alunos['ex_g'].tolist()), key=f"sag_{st.session_state.al_form_id}")
                    al_edit = df_alunos[df_alunos['aluno'] == sel_a_g.split(" - ")[0]].iloc[0]
                with st.form(key=f"f_al_{st.session_state.al_form_id}"):
                    reg = st.text_input("Registro", value=al_edit['registro'] if al_edit is not None else "")
                    nom = st.text_input("Nome", value=al_edit['aluno'] if al_edit is not None else "")
                    tur = st.text_input("Turma", value=al_edit['turma'] if al_edit is not None else "")
                    nec = st.text_area("Condição", value=al_edit['necessidades'] if al_edit is not None else "")
                    nas = st.text_input("Nascimento", value=al_edit['data_nascimento'] if al_edit is not None else "")
                    obs = st.text_area("Observações", value=al_edit['observacoes_gerais'] if al_edit is not None else "")
                    ft_p = st.file_uploader("Foto Perfil")
                    c1, c2 = st.columns(2)
                    if c1.form_submit_button("Salvar Perfil"):
                        p_path = al_edit['foto_path'] if al_edit is not None else ""
                        if ft_p:
                            p_path = f"perfil_{reg}.png"
                            supabase.storage.from_("fotos_perfil").upload(p_path, ft_p.getvalue(), {"upsert": "true"})
                        supabase.table("estudantes").upsert({"registro": reg, "aluno": nom, "turma": tur, "necessidades": nec, "data_nascimento": nas, "observacoes_gerais": obs, "foto_path": p_path}).execute()
                        st.toast("✅ Aluno salvo!"); st.session_state.al_form_id += 1; time.sleep(1); st.rerun()
                    if modo_a == "Editar/Excluir" and c2.form_submit_button("❌ EXCLUIR ALUNO"):
                        if al_edit['foto_path']: supabase.storage.from_("fotos_perfil").remove([al_edit['foto_path']])
                        supabase.table("estudantes").delete().eq("registro", al_edit['registro']).execute()
                        st.session_state.al_form_id += 1; time.sleep(1); st.rerun()

            with abas[3]: # GESTÃO DE PROFESSORES
                st.subheader("Gerenciar Professores")
                if "p_form_id" not in st.session_state: st.session_state.p_form_id = 0
                modo_p = st.radio("Ação Professor:", ["Novo", "Editar/Excluir"], horizontal=True, key=f"mp_{st.session_state.p_form_id}")
                perfis_op = ["professor", "paee", "direcao", "coordenador"]
                if st.session_state.u_perfil == "gestao": perfis_op.append("gestao")
                if modo_p == "Novo":
                    with st.form(key=f"fn_{st.session_state.p_form_id}"):
                        nr, nn = st.text_input("RF"), st.text_input("Nome"); np = st.selectbox("Perfil", perfis_op)
                        if st.form_submit_button("Cadastrar"):
                            if not df_prof[df_prof['rf'] == nr].empty: st.error("RF já existe!"); st.stop()
                            supabase.table("professores").insert({"rf": nr, "nome": nn, "perfil": np}).execute()
                            st.toast("✅ Cadastrado!"); st.session_state.p_form_id += 1; time.sleep(1); st.rerun()
                else:
                    l_edit = sorted(df_prof['nome'].tolist()) if st.session_state.u_perfil == "gestao" else sorted(df_prof[df_prof['perfil'] != 'gestao']['nome'].tolist())
                    psn = st.selectbox("Selecionar:", l_edit)
                    psd = df_prof[df_prof['nome'] == psn].iloc[0]
                    with st.form(f"fe_{psn}"):
                        en, ep = st.text_input("Nome", value=psd['nome']), st.selectbox("Perfil", perfis_op, index=perfis_op.index(psd['perfil']) if psd['perfil'] in perfis_op else 0)
                        c_b1, c_b2 = st.columns(2)
                        if c_b1.form_submit_button("Atualizar"): supabase.table("professores").update({"nome": en, "perfil": ep}).eq("rf", psd['rf']).execute(); st.rerun()
                        if c_b2.form_submit_button("Excluir"):
                            supabase.table("professores").delete().eq("rf", psd['rf']).execute()
                            supabase.table("credenciais").delete().eq("rf", psd['rf']).execute(); st.rerun()

            with abas[4]: # SEGURANÇA E RESET
                st.subheader("Segurança e Monitoramento")
                col1, col2 = st.columns(2)
                with col1:
                    pr = st.selectbox("Resetar Professor:", sorted(df_prof[df_prof['perfil'] != 'gestao']['nome'].tolist()) if st.session_state.u_perfil != "gestao" else sorted(df_prof['nome'].tolist()))
                    if st.button("Resetar Senha"): supabase.table("credenciais").delete().eq("rf", df_prof[df_prof['nome'] == pr].iloc[0]['rf']).execute(); st.warning("Resetado.")
                    if st.button("📊 Monitoramento Excel"):
                        all_r = pd.DataFrame(supabase.table("relatorios").select("*").execute().data)
                        if not all_r.empty:
                            m = all_r.merge(df_prof[['rf', 'nome']], left_on='rf_professor', right_on='rf', how='left').merge(df_alunos[['registro', 'aluno', 'turma']], left_on='registro_aluno', right_on='registro', how='left')
                            m = m[['nome_y', 'aluno', 'turma', 'data', 'bimestre']]; m.columns = ['Professor', 'Aluno', 'Turma', 'Data', 'Bimestre']
                            out = BytesIO(); m.to_excel(out, index=False); st.download_button("Download Excel", out.getvalue(), "monitor.xlsx")
                with col2:
                    if st.session_state.u_perfil == "gestao":
                        st.error("🚨 RESET TOTAL")
                        if st.button("🚨 ZERAR TUDO"): st.session_state.conf_res = True
                        if st.session_state.get("conf_res") and st.button("CONFIRMAR APAGAMENTO"):
                            for t in ["relatorios", "logs", "credenciais", "estudantes", "professores"]:
                                col_pk = "id" if t in ["relatorios", "logs"] else ("rf" if t in ["professores", "credenciais"] else "registro")
                                supabase.table(t).delete().neq(col_pk, "xxx").execute()
                            st.session_state.logged_in = False; st.rerun()